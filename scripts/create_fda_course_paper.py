from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path('/Users/jao/Desktop/FDA-Inspections/FDA_Inspection_Risk_Model_Paper.docx')


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, align=None):
    cell.text = ''
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)


def add_note_marker(paragraph, number):
    run = paragraph.add_run(str(number))
    run.font.superscript = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)


def add_paragraph_with_note(text_before, note_number, text_after='', style=None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text_before)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    add_note_marker(paragraph, note_number)
    if text_after:
        run2 = paragraph.add_run(text_after)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
    return paragraph


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles['Normal'].font.name = 'Times New Roman'
styles['Normal'].font.size = Pt(12)

for style_name in ['Title', 'Subtitle', 'Heading 1', 'Heading 2']:
    styles[style_name].font.name = 'Times New Roman'

styles['Title'].font.size = Pt(16)
styles['Title'].font.bold = True
styles['Subtitle'].font.size = Pt(11)
styles['Subtitle'].font.italic = True
styles['Heading 1'].font.size = Pt(13)
styles['Heading 1'].font.bold = True
styles['Heading 2'].font.size = Pt(12)
styles['Heading 2'].font.bold = True

p = doc.add_paragraph(style='Title')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('FDA Inspection Classification Risk Modeling')
r.font.name = 'Times New Roman'
r.font.size = Pt(16)
r.bold = True

p = doc.add_paragraph(style='Subtitle')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Prepared for FDA course discussion based on the project notebook and browser deployment')
r.font.name = 'Times New Roman'
r.font.size = Pt(11)
r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Prepared April 25, 2026')
r.font.name = 'Times New Roman'
r.font.size = Pt(11)

# Introduction
h = doc.add_heading('Introduction', level=1)
for run in h.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.bold = True

intro = (
    'This paper describes a browser-based machine learning project that estimates FDA inspection outcomes from '
    'minimal case information. The project uses the FDA inspection classification dataset, a LightGBM multiclass '
    'model, and a static web interface that returns probabilities for No Action Indicated (NAI), Voluntary Action '
    'Indicated (VAI), and Official Action Indicated (OAI). The goal is not to replace FDA judgment, but to show how '
    'public inspection data can be converted into a transparent predictive workflow that is easy to inspect, test, '
    'and deploy.'
)
doc.add_paragraph(intro)

# Regulatory framework
h = doc.add_heading('Regulatory Framework', level=1)
for run in h.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.bold = True

para = doc.add_paragraph()
para.add_run(
    'The legal foundation for FDA inspections comes from the Federal Food, Drug, and Cosmetic Act, which authorizes '
    'FDA to inspect regulated facilities and gather information needed to determine compliance with applicable law. '
)
add_note_marker(para, 1)
para.add_run(
    ' In the agency\'s own explanation of the inspection database, FDA states that inspections are used to verify '
    'compliance with the laws it administers, to support surveillance, and to follow up on specific risk factors. '
    'The inspection database is risk based, is updated weekly, and does not include every inspection the agency '
    'conducts.'
)
add_note_marker(para, 2)

para = doc.add_paragraph()
para.add_run(
    'The database organizes results by project area within an inspection and assigns one of three final '
    'classifications. NAI means the firm is in an acceptable state of compliance; VAI means objectionable conditions '
    'were found, but voluntary correction is deemed sufficient; and OAI means the firm is in an unacceptable state '
    'of compliance. These labels are therefore best understood as administrative classifications, not as criminal or '
    'civil adjudications. They also reflect only the final classification for each project area and do not capture all '
    'inspection activity across the FDA system.'
)
add_note_marker(para, 3)

para = doc.add_paragraph()
para.add_run(
    'That regulatory structure matters for modeling because the target variable is inherently imbalanced and only '
    'partially observable in public data. A predictive model built from the dashboard necessarily learns from the '
    'subset of inspections that FDA publishes, not from the full universe of inspection activity. For that reason, '
    'the model should be read as a statistical summary of public outcomes rather than a compliance determination.'
)

# Data and model
h = doc.add_heading('Data and Model Design', level=1)
for run in h.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.bold = True

para = doc.add_paragraph()
para.add_run(
    'The project uses the FDA Inspection Dashboard dataset as its primary source and maps the browser interface to a '
    'small set of ex ante inputs. The notebook engineerings the remaining features from the user\'s selections. '
    'Those features include state, zip code, country or area, fiscal year, project area, product type, inspection '
    'history counts, a presidential administration mapping, a party indicator, years since administration start, '
    'and frequency-based product rankings. The model also converts country and state information into geographic '
    'regions so that sparse categorical inputs can be represented in a consistent form.'
)
add_note_marker(para, 4)

para = doc.add_paragraph()
para.add_run(
    'The model is a LightGBM multiclass classifier that preserves the fitted tree ensemble for browser-side '
    'inference. The notebook splits the data by time to respect the chronology of the inspection records: training '
    'uses inspections through fiscal year 2023, validation uses fiscal year 2024, and the final test set consists of '
    'fiscal years 2025 and 2026. This temporal split is important because it avoids training on future outcomes and '
    'better approximates how the model would behave prospectively.'
)

para = doc.add_paragraph()
para.add_run(
    'The project chooses a sparse, interpretable feature set by design. Instead of relying on post-inspection or '
    'other leakage-prone variables, it uses only information that would plausibly be known at the time of prediction. '
    'That choice makes the browser app more realistic and also easier to explain in a course setting.'
)

# Model summary table
h = doc.add_heading('Model Summary', level=1)
for run in h.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.bold = True

summary_table = doc.add_table(rows=1, cols=2)
summary_table.style = 'Table Grid'
summary_table.rows[0].cells[0].text = 'Metric'
summary_table.rows[0].cells[1].text = 'Value'
for cell in summary_table.rows[0].cells:
    set_cell_shading(cell, 'D9EAF7')
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

summary_rows = [
    ('Training window', 'Fiscal years 2009 to 2023'),
    ('Validation window', 'Fiscal year 2024'),
    ('Test window', 'Fiscal years 2025 to 2026'),
    ('Model type', 'LightGBM multiclass classifier'),
    ('Output', 'NAI, VAI, and OAI probability profile'),
]
for metric, value in summary_rows:
    row = summary_table.add_row().cells
    set_cell_text(row[0], metric)
    set_cell_text(row[1], value)

# Results
h = doc.add_heading('Results', level=1)
for run in h.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.bold = True

para = doc.add_paragraph()
para.add_run(
    'On the held-out test set, the model achieved 73.24 percent accuracy and 42.29 percent macro-F1. Accuracy is '
    'helpful as a broad summary, but macro-F1 is more informative here because the classes are imbalanced and the '
    'rare OAI class is much harder to predict. The confusion matrix shows that the model is strongest on NAI and '
    'substantially weaker on VAI and especially OAI.'
)

results_table = doc.add_table(rows=1, cols=4)
results_table.style = 'Table Grid'
headers = ['Class', 'Recall', 'Precision', 'Support']
for i, header in enumerate(headers):
    results_table.rows[0].cells[i].text = header
    set_cell_shading(results_table.rows[0].cells[i], 'D9EAF7')
    for paragraph in results_table.rows[0].cells[i].paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

result_rows = [
    ('NAI', '91.81%', '76.81%', '17,706'),
    ('VAI', '31.08%', '54.17%', '6,709'),
    ('OAI', '2.02%', '26.53%', '645'),
]
for row_data in result_rows:
    row = results_table.add_row().cells
    for idx, value in enumerate(row_data):
        set_cell_text(row[idx], value, align=WD_ALIGN_PARAGRAPH.CENTER)

para = doc.add_paragraph()
para.add_run(
    'The class-level behavior is consistent with the distribution chart in the notebook. Across fiscal years, NAI '
    'remains dominant, VAI occupies a smaller but still substantial share, and OAI remains a small minority class. '
    'That pattern helps explain why the classifier performs much better on NAI than on OAI. In practical terms, the '
    'model learns the majority pattern well, but it does not yet separate the rare enforcement-heavy outcomes with '
    'high reliability.'
)

# Confusion matrix table
h = doc.add_heading('Confusion Matrix', level=2)
for run in h.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True

cm_table = doc.add_table(rows=4, cols=4)
cm_table.style = 'Table Grid'
cm_headers = ['', 'Predicted NAI', 'Predicted VAI', 'Predicted OAI']
for i, header in enumerate(cm_headers):
    cm_table.rows[0].cells[i].text = header
    set_cell_shading(cm_table.rows[0].cells[i], 'D9EAF7')
    for paragraph in cm_table.rows[0].cells[i].paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

cm_rows = [
    ('Actual NAI', '16,255', '1,443', '8'),
    ('Actual VAI', '4,596', '2,085', '28'),
    ('Actual OAI', '311', '321', '13'),
]
for row_idx, row_data in enumerate(cm_rows, start=1):
    for col_idx, value in enumerate(row_data):
        set_cell_text(cm_table.rows[row_idx].cells[col_idx], value, align=WD_ALIGN_PARAGRAPH.CENTER)

para = doc.add_paragraph()
para.add_run(
    'The confusion matrix shows that most errors come from the model moving observations into adjacent classes. '
    'For example, many actual VAI inspections are predicted as NAI, and nearly all actual OAI inspections are '
    'predicted as either NAI or VAI. That result is not surprising in a dataset where OAI is rare and where the '
    'browser app only uses ex ante variables.'
)

# Limitations and conclusion
h = doc.add_heading('Limitations and Conclusion', level=1)
for run in h.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.bold = True

para = doc.add_paragraph()
para.add_run(
    'The model is useful for classroom analysis and for illustrating how public FDA data can be turned into a '
    'predictive interface, but it should not be treated as a regulatory instrument. The public database excludes '
    'some inspections, groups results by project area, and reports final classifications rather than every step of '
    'the inspection process. Those design choices limit completeness, and they also mean the model is learning from '
    'a curated administrative record rather than the entire universe of FDA oversight activity.'
)

para = doc.add_paragraph()
para.add_run(
    'Even with those limits, the project demonstrates a coherent regulatory analytics workflow. It links the '
    'statutory inspection authority, the FDA\'s public inspection database, a feature-engineered tabular model, and '
    'a browser deployment that returns class probabilities rather than a hard enforcement conclusion. For an FDA '
    'course, that combination is useful because it shows how administrative law, data constraints, and machine '
    'learning design intersect in a single applied project.'
)

# Notes / authorities
h = doc.add_heading('Authorities Cited', level=1)
for run in h.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.bold = True

notes = [
    '21 U.S.C. § 374 (2024).',
    'U.S. Food & Drug Admin., Inspections Database Frequently Asked Questions, https://www.fda.gov/inspections-compliance-enforcement-and-criminal-investigations/inspection-references/inspections-database-frequently-asked-questions (last visited Apr. 25, 2026).',
    'U.S. Food & Drug Admin., United States Food and Drug Administration\'s Data Dashboard: Inspections, https://datadashboard.fda.gov/oii/cd/inspections.htm (last visited Apr. 25, 2026).',
    'FDA inspection classification labels and dataset caveats are discussed in the FDA Inspections Database Frequently Asked Questions and the FDA Data Dashboard page cited above.',
]
for idx, note in enumerate(notes, start=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    marker = paragraph.add_run(f'{idx}. ')
    marker.font.name = 'Times New Roman'
    marker.font.size = Pt(11)
    body = paragraph.add_run(note)
    body.font.name = 'Times New Roman'
    body.font.size = Pt(11)

# Add hidden page break style section for clean ending
section = doc.sections[0]
section.start_type = WD_SECTION_START.NEW_PAGE

OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT_PATH)
print(f'Wrote {OUTPUT_PATH}')
